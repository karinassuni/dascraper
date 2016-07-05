import docx
import datetime
import json

WORD_DOC = docx.Document("ClubMeetingsSpring2016.docx")

def main():
    START_INDEX = 3
    NAME_INDEX, DAYS_INDEX, DATES_INDEX, TIME_INDEX, LOCATION_INDEX = range(1, 6)
    clubs = []
    print "Parsing the club spreadsheet..."
    for table in WORD_DOC.tables:
        for row in table.rows[START_INDEX:]:
            club_is_active = bool(row.cells[DATES_INDEX].text.strip())
            if club_is_active:
                club = {
                    "name": row.cells[NAME_INDEX].text.strip(),
                    "days": extract_days(row.cells[DAYS_INDEX].text),
                    "dates": extract_dates(row.cells[DATES_INDEX].text),
                    "start_time": clean_time(row.cells[TIME_INDEX].text.split(" - ")[0]),
                    "location": row.cells[LOCATION_INDEX].text.strip()
                }
                try:
                    club["end_time"] = clean_time(row.cells[TIME_INDEX].text.split(" - ")[1])
                except IndexError:
                    club["end_time"] = ''
                clubs.append(club)
    with open("clubs.json", 'w') as outfile:
        json.dump(clubs, outfile)
    print "Finished!"

def extract_days(days):
    # Keep only alphanumeric characters, and whitespace for splitting the string into list items
    days_array = ''.join(char for char in days if char.isalnum() or char.isspace()).split()
    for i, day in enumerate(days_array):
        days_array[i] = days_array[i][0:3].capitalize()

    return days_array

def extract_dates(dates):
    """ input:
        4/8, 15, 22, 29, 5/6, 13, 20, 27, 6/3, 10, 17
        output:
        4/8/16, 4/15/16, 4/22/16, 4/29/16
    """

    month = ''
    WORD_DOC_YEAR = WORD_DOC.tables[0].cell(0,3).text.split()[0]
    dates_array = dates.split(", ")
    # Since we need to modify each list element as we loop, we need to access indicies
    for i, date in enumerate(dates_array):
        date_has_month = bool('/' in date)
        if date_has_month:
            month = date[:date.index('/')]
            # strip month
            date = date[date.index('/')+1:]
        try:
            dates_array[i] = datetime.date(int(WORD_DOC_YEAR), int(month), int(date)).isoformat()
        except ValueError:
            continue

    return dates_array

def clean_time(time):
    time = time.lower()
    if 'a' in time:
        time = time.split()[0] + "AM"
    elif 'p' in time:
        time = time.split()[0] + "PM"
    return datetime.datetime.strptime(time, "%I:%M%p").time().isoformat()

if __name__ == "__main__":
    main()
