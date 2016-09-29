import datetime
import docx
import json
import logging
import os
from dascraper.utility import clean_time


# Path arguments in os.path are relative to the present working directory
# (the directory from where the module is called), so use the unchanging
# absolute path to the directory containing this file via `__file__`
WORD_DOC = docx.Document(os.path.join(
    os.path.dirname(__file__), os.path.relpath("spreadsheet.docx")))


def parse():
    logging.info("Parsing the club spreadsheet...")

    COLUMNS = ('', "name", "meetingDays", "meetings", "time", "location", '')
    FIRST_ROW = 3
    clubs = {}

    for table in WORD_DOC.tables:
        for row in table.rows[FIRST_ROW:]:
            row_cells = row.cells
            club_is_active = bool(
                row_cells[COLUMNS.index("meetings")]
                .text.strip() != ''
            )
            if club_is_active:
                try:
                    club = clean({
                        COLUMNS[i]: cell.text
                        for i, cell in enumerate(row_cells)
                    })
                except ValueError:
                    continue
                key = club.pop("name")
                clubs[key] = club

    logging.info("Finished parsing the club spreadsheet")
    return clubs


def clean(club):
    for field, value in club.items():
        club[field] = value.strip()

    club["meetingDays"] = split_days(club["meetingDays"])

    try:
        start_time, end_time = club["time"].split('-')
    # Not all club meetings have end times; if not, equate to start_time to
    # produce a valid DateTime for here and for clients
    except ValueError:
        start_time = end_time = club["time"]

    start_time, end_time = tuple(
        clean_time.meridiem(t)
        for t in (start_time, end_time)
    )

    meeting_intervals = []
    for meeting_date in split_dates(club["meetings"]):
        # 1 Date + 2 times = 2 DateTimes
        start, end = tuple(
            clean_time.to_utc_datetime(meeting_date, time).isoformat()
            for time in (start_time, end_time)
        )
        meeting_intervals.append(start + '/' + end)
    club["meetings"] = meeting_intervals

    # start_time and end_time found; raw "time" no longer needed
    club.pop("time", None)

    # Remove '' key generated from the blank COLUMNS
    club.pop('', None)

    club["facebookUrl"] = ''
    club["fromSpreadsheet"] = True

    return club


def split_days(days):
    """
    Given a string of days with arbitrary separators, return an array of ISO-
    friendly day strings
    >>> split_days("Thursdays / Fridays")
        ["Thu", "Fri"]
    """

    # Leave only letters and spaces, so that split() works consistently
    days_list = ''.join(
        char
        for char in days
        if char.isalpha()
        or char.isspace()
    ).split()

    day_switch = {
        "Su" : "SUNDAY",
        "Mo" : "MONDAY",
        "Tu" : "TUESDAY",
        "We" : "WEDNESDAY",
        "Th" : "THURSDAY",
        "Fr" : "FRIDAY",
        "Sa" : "SATURDAY"
    }

    return [
        day_switch[day]
        for day in [d[:2] for d in days_list]
    ]



def split_dates(dates):
    """
    Given a string of dates with arbitrary separators, return an array of ISO
    dates
    >>> split_dates("4/8, 15, 22, 29, 620")
        ["2016-04-08", "2016-04-15", "2016-04-22", "2016-04-29"]
    """

    # Cell text example: "2016 SPRING CLUB MEETING SCHEDULE"
    WORD_DOC_YEAR = int(WORD_DOC.tables[0].cell(0, 3).text.split()[0])

    # Remove all whitespace from dates string, leaving only commas for splitting
    dates_list = ''.join(dates.split()).split(',')

    invalid_dates = []

    # Since we're modifying the list in place, we need the current index
    for i, date in enumerate(dates_list):
        date_has_month = bool('/' in date)
        if date_has_month:
            # Based on the '/' month separator, record the month and date
            month = int(date[:date.index('/')])
            day = int(date[date.index('/') + 1:])
        else:
            day = int(date)

        try:
            dates_list[i] = datetime.date(WORD_DOC_YEAR, month, day).isoformat()
        except ValueError:
            logging.debug("Invalid date in spreadsheet: \"{}\"".format(date))
            invalid_dates.append(date)
            continue

    return [date for date in dates_list if date not in invalid_dates]


def main():
    with open("clubs_without_descriptions.json", 'w') as o:
        json.dump(parse(), o, sort_keys=True, indent=4, separators=(',', ': '))


if __name__ == "__main__":
    main()
