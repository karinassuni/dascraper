import dascraper.cleantime as cleantime
import datetime
import docx
import json
import logging
import os


# Path arguments in os.path are relative to the present working directory
# (the directory from where the module is called), so use the unchanging
# absolute path to the directory containing this file via `__file__`
WORD_DOC = docx.Document(os.path.join(
    os.path.dirname(__file__), os.path.relpath("res/ClubMeetingsSpring2016.docx")))


def parse():
    RAW_FIELDS = ('', "name", "days", "dates", "time", "location", '')
    FIRST_ROW = 3
    clubs = []

    logging.debug("Parsing the club spreadsheet...")

    for table in WORD_DOC.tables:
        for row in table.rows[FIRST_ROW:]:
            row_cells = row.cells
            club_is_active = bool(row_cells[(RAW_FIELDS.index("dates"))].text.strip() != '')
            if club_is_active:
                club = clean({
                    RAW_FIELDS[i]: cell.text
                    for i, cell in enumerate(row_cells)
                })
                clubs.append(club)

    logging.debug("Finished parsing the club spreadsheet")

    return clubs


def clean(club):
    clean_club = club
    clean_club["name"] = club["name"].strip()
    clean_club["days"] = extract_days(club["days"])
    clean_club["dates"] = extract_dates(club["dates"])
    clean_club["start_time"] = cleantime.iso(club["time"].split(" - ")[0])

    try:
        clean_club["end_time"] = cleantime.iso(club["time"].split(" - ")[1])
    except IndexError:
        clean_club["end_time"] = ''

    clean_club["location"] = club["location"].strip()

    # start_time and end_time found; raw "time" no longer needed
    clean_club.pop("time", None)

    # Remove '' key generated from the blank raw fields
    clean_club.pop('', None)

    return clean_club


def extract_days(days):
    """
    Given a string of days with arbitrary separators, return an array of ISO-
    friendly day strings
    >>> extract_days("Thursdays / Fridays")
        ["Thu", "Fri"]
    """

    # Leave only letters and spaces, so that split() works consistently
    clean_days = ''.join(c for c in days if c.isalpha() or c.isspace()).split()

    for i, day in enumerate(clean_days):
        clean_days[i] = day[0:3].capitalize()

    return clean_days


def extract_dates(dates):
    """
    Given a string of dates with arbitrary separators, return an array of ISO dates
    >>> extract_dates("4/8, 15, 22, 29")
        ["2016-04-08", "2016-04-15", "2016-04-22", "2016-04-29"]
    """

    WORD_DOC_YEAR = int(WORD_DOC.tables[0].cell(0, 3).text.split()[0])

    # Remove all whitespace from dates string, leaving only commas for splitting
    clean_dates = ''.join(dates.split()).split(',')

    month = 0

    # Since we're modifying the array while looping, we need the index
    for i, date in enumerate(clean_dates):
        date_has_month = bool('/' in date)
        if date_has_month:
            month = int(date[:date.index('/')])
            # Strip the date's month, for uniformity outside the if block
            date = date[date.index('/') + 1:]
        try:
            clean_dates[i] = datetime.date(
                WORD_DOC_YEAR, month, int(date)).isoformat()
        except ValueError:
            logging.exception("Invalid date in spreadsheet: \"{}\"".format(date))
            continue

    return clean_dates


def main():
    with open("clubs.json", 'w') as o:
        json.dump(parse(), o, sort_keys=True)


if __name__ == "__main__":
    main()
